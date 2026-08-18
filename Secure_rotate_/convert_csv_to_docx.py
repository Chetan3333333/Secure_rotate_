import os
import pandas as pd
from docx import Document

DATASETS_DIR = "datasets"

def csv_to_docx(csv_file_path, docx_file_path):
    # Read the CSV file
    try:
        df = pd.read_csv(csv_file_path)
    except Exception as e:
        print(f"Error reading {csv_file_path}: {e}")
        return

    # Create a new Word document
    doc = Document()
    doc.add_heading(os.path.basename(csv_file_path), level=1)

    # Add a table to the document
    # Number of rows is dataframe rows + 1 (for header), cols is dataframe cols
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = 'Table Grid'

    # Add the header row
    for j, col_name in enumerate(df.columns):
        table.cell(0, j).text = str(col_name)

    # Add the data rows
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            # i+1 because 0 is header
            table.cell(i + 1, j).text = str(df.iat[i, j])

    # Save the document
    doc.save(docx_file_path)
    print(f"Successfully converted {csv_file_path} to {docx_file_path}")

def main():
    if not os.path.exists(DATASETS_DIR):
        print(f"Directory '{DATASETS_DIR}' does not exist.")
        return

    csv_files = [f for f in os.listdir(DATASETS_DIR) if f.endswith(".csv")]
    
    if not csv_files:
        print(f"No CSV files found in '{DATASETS_DIR}'.")
        return
        
    for csv_file in csv_files:
        csv_path = os.path.join(DATASETS_DIR, csv_file)
        docx_name = csv_file.replace(".csv", ".docx")
        docx_path = os.path.join(DATASETS_DIR, docx_name)
        
        csv_to_docx(csv_path, docx_path)

if __name__ == "__main__":
    main()
